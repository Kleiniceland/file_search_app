import os
import sys
import re

# Check for python-docx
try:
    import docx
except ImportError:
    print("=" * 50)
    print("缺少 python-docx 库。")
    print("请在命令行中执行以下命令安装：")
    print("  pip install python-docx")
    print("然后重新运行此脚本。")
    print("=" * 50)
    sys.exit(1)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def parse_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('> '):
            # Blockquote
            p = doc.add_paragraph()
            p.add_run(line[2:]).italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.startswith('---'):
            # Horizontal rule
            doc.add_paragraph('─' * 80)
            
        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add a light gray background
            p.paragraph_format.left_indent = Inches(0.3)
            # Note: python-docx doesn't easily support background color on paragraph, 
            # we just format the text
            
        # Table
        elif line.startswith('|'):
            # Check if it's a table header
            if i + 1 < len(lines) and lines[i+1].startswith('|') and '---' in lines[i+1]:
                # Parse table
                header_cells = [c.strip() for c in line.split('|')[1:-1]]
                i += 2 # skip separator line
                data_rows = []
                while i < len(lines) and lines[i].startswith('|'):
                    row_cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                    data_rows.append(row_cells)
                    i += 1
                
                # Create table
                table = doc.add_table(rows=1+len(data_rows), cols=len(header_cells))
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for j, text in enumerate(header_cells):
                    hdr_cells[j].text = text
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Data
                for row_idx, row_data in enumerate(data_rows):
                    row_cells = table.rows[row_idx+1].cells
                    for col_idx, text in enumerate(row_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = text
                continue
            else:
                # Just a line starting with | but not a table, treat as paragraph
                pass
                
        # List items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            
        else:
            # Normal paragraph
            # Check for bold or italic inline formatting
            clean_text = line
            # Simple markdown processing for bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            if len(parts) > 1:
                p = doc.add_paragraph()
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        # Handle links [text](url)
                        link_match = re.match(r'\[.*?\]\(.*?\)', part)
                        if link_match:
                            run = p.add_run(part)
                            run.italic = True
                        else:
                            p.add_run(part)
            else:
                # Handle links [text](url)
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', line)
                if link_match:
                    p = doc.add_paragraph()
                    run = p.add_run(link_match.group(1))
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
                else:
                    doc.add_paragraph(line)
        
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    
    md_file = os.path.join(project_root, "docs", "PROJECT_FINAL_REPORT.md")
    docx_file = os.path.join(project_root, "docs", "PROJECT_FINAL_REPORT.docx")
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found")
        sys.exit(1)
        
    parse_md_to_docx(md_file, docx_file)
